"""
文档操作工具集
这些函数会被 LangChain Agent 调用
"""
from langchain.tools import tool
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import Optional, List, Dict
import re
import json
from PIL import Image
import base64
from io import BytesIO

@tool
def analyze_template_structure(file_path: str) -> str:
    """
    深度分析模板文档的样式定义
    
    这个工具会从模板的样式定义（而不是段落内容）中提取样式信息：
    - 所有标题样式（Heading 1-6 或中文样式如"标题 1"）
    - 正文样式（Normal、Body Text 等）
    - 每个样式的详细属性（字体、大小、粗体、斜体、对齐、颜色等）
    
    Args:
        file_path: 模板文件路径
        
    Returns:
        模板样式的详细分析结果（JSON格式 + 文字摘要）
    """
    try:
        doc = Document(file_path)
        
        # 辅助函数：将 EMU 转换为 Pt
        def emu_to_pt(emu):
            if emu is None:
                return None
            return round(emu / 12700, 1)  # 1 Pt = 12700 EMU
        
        # 辅助函数：获取对齐方式名称
        def get_alignment_name(alignment):
            if alignment is None:
                return "LEFT"
            alignment_map = {
                0: "LEFT",
                1: "CENTER", 
                2: "RIGHT",
                3: "JUSTIFY"
            }
            return alignment_map.get(int(alignment), "LEFT")
        
        # 1. 从文档样式定义中提取样式
        heading_styles = []
        body_styles = []
        
        # 定义要查找的样式名称
        heading_names = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6',
                         '标题 1', '标题 2', '标题 3', '标题 4', '标题', 'Title', 'Subtitle']
        body_names = ['Normal', 'Body Text', 'Body Text First Indent', '正文', '正文首行缩进', 
                      '正文 2', 'Body', 'Paragraph']
        
        # 收集文档中所有可用的样式
        available_styles = {}
        for style in doc.styles:
            if hasattr(style, 'name') and style.name:
                available_styles[style.name] = style
        
        # 提取标题样式
        for style_name in heading_names:
            if style_name in available_styles:
                style = available_styles[style_name]
                style_info = {
                    "name": style_name,
                    "type": "heading",
                    "font_name": None,
                    "font_size_pt": None,
                    "bold": None,
                    "italic": None,
                    "alignment": None,
                    "color": None
                }
                
                # 提取字体属性
                if hasattr(style, 'font') and style.font:
                    style_info["font_name"] = style.font.name
                    style_info["font_size_pt"] = emu_to_pt(style.font.size)
                    style_info["bold"] = style.font.bold
                    style_info["italic"] = style.font.italic
                    if style.font.color and style.font.color.rgb:
                        style_info["color"] = str(style.font.color.rgb)
                
                # 提取段落属性
                if hasattr(style, 'paragraph_format') and style.paragraph_format:
                    style_info["alignment"] = get_alignment_name(style.paragraph_format.alignment)
                
                heading_styles.append(style_info)
        
        # 提取正文样式
        for style_name in body_names:
            if style_name in available_styles:
                style = available_styles[style_name]
                style_info = {
                    "name": style_name,
                    "type": "body",
                    "font_name": None,
                    "font_size_pt": None,
                    "bold": None,
                    "italic": None,
                    "alignment": None,
                    "first_line_indent": None
                }
                
                if hasattr(style, 'font') and style.font:
                    style_info["font_name"] = style.font.name
                    style_info["font_size_pt"] = emu_to_pt(style.font.size)
                    style_info["bold"] = style.font.bold
                    style_info["italic"] = style.font.italic
                
                if hasattr(style, 'paragraph_format') and style.paragraph_format:
                    style_info["alignment"] = get_alignment_name(style.paragraph_format.alignment)
                    if style.paragraph_format.first_line_indent:
                        style_info["first_line_indent"] = emu_to_pt(style.paragraph_format.first_line_indent)
                
                body_styles.append(style_info)
        
        # 2. 如果样式定义为空，从段落中推断样式
        if not heading_styles and not body_styles:
            # 从段落中收集实际使用的样式
            used_styles = set()
            for para in doc.paragraphs:
                if para.text.strip():
                    used_styles.add(para.style.name)
            
            for style_name in used_styles:
                if style_name in available_styles:
                    style = available_styles[style_name]
                    style_info = {
                        "name": style_name,
                        "type": "heading" if "Heading" in style_name or "标题" in style_name else "body",
                        "font_name": getattr(style.font, 'name', None) if hasattr(style, 'font') else None,
                        "font_size_pt": emu_to_pt(getattr(style.font, 'size', None)) if hasattr(style, 'font') else None,
                        "bold": getattr(style.font, 'bold', None) if hasattr(style, 'font') else None,
                    }
                    if style_info["type"] == "heading":
                        heading_styles.append(style_info)
                    else:
                        body_styles.append(style_info)
        
        # 3. 统计文档内容
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)
        image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.target_ref)
        
        # 4. 构建结果
        result = {
            "heading_styles": heading_styles,
            "body_styles": body_styles,
            "document_stats": {
                "total_paragraphs": total_paragraphs,
                "total_tables": total_tables,
                "total_images": image_count
            }
        }
        
        # 5. 生成文字摘要
        summary = f"[OK] 模板分析完成\n\n"
        summary += f"[统计] 文档统计：{total_paragraphs}个段落，{total_tables}个表格，{image_count}张图片\n\n"
        
        if heading_styles:
            summary += "[标题样式]：\n"
            for hs in heading_styles[:5]:
                size_info = f"{hs['font_size_pt']}pt" if hs['font_size_pt'] else "默认"
                bold_info = "加粗" if hs['bold'] else ""
                summary += f"  - {hs['name']}: {size_info} {bold_info}\n"
        
        if body_styles:
            summary += "\n[正文样式]：\n"
            for bs in body_styles[:3]:
                size_info = f"{bs['font_size_pt']}pt" if bs['font_size_pt'] else "默认"
                summary += f"  - {bs['name']}: {size_info}\n"
        
        summary += f"\n---\n【JSON数据】\n{json.dumps(result, ensure_ascii=False, indent=2)}"
        
        return summary
    
    except Exception as e:
        return f"分析模板失败：{str(e)}"

@tool
def get_template_styles(file_path: str) -> str:
    """
    获取模板中可用的样式列表（供前端使用）
    
    返回结构化的JSON，包含所有可用的标题样式和正文样式，
    前端可以用这个列表让用户选择每段内容应该使用什么样式。
    
    Args:
        file_path: 模板文件路径
        
    Returns:
        JSON格式的样式列表
    """
    try:
        doc = Document(file_path)
        
        def emu_to_pt(emu):
            if emu is None:
                return None
            return round(emu / 12700, 1)
        
        def get_alignment_name(alignment):
            if alignment is None:
                return "LEFT"
            alignment_map = {0: "LEFT", 1: "CENTER", 2: "RIGHT", 3: "JUSTIFY"}
            return alignment_map.get(int(alignment), "LEFT")
        
        # 收集所有样式
        styles_result = {
            "headings": [],
            "body": [],
            "all_styles": []
        }
        
        heading_keywords = ['Heading', '标题', 'Title', 'Subtitle']
        body_keywords = ['Normal', 'Body', '正文', 'Paragraph']
        
        for style in doc.styles:
            if not hasattr(style, 'name') or not style.name:
                continue
            
            style_info = {
                "name": style.name,
                "font_size_pt": None,
                "bold": None,
                "alignment": None
            }
            
            if hasattr(style, 'font') and style.font:
                style_info["font_size_pt"] = emu_to_pt(style.font.size)
                style_info["bold"] = style.font.bold
            
            if hasattr(style, 'paragraph_format') and style.paragraph_format:
                style_info["alignment"] = get_alignment_name(style.paragraph_format.alignment)
            
            # 分类
            is_heading = any(kw in style.name for kw in heading_keywords)
            is_body = any(kw in style.name for kw in body_keywords)
            
            if is_heading:
                styles_result["headings"].append(style_info)
            elif is_body:
                styles_result["body"].append(style_info)
            
            styles_result["all_styles"].append(style.name)
        
        return json.dumps(styles_result, ensure_ascii=False, indent=2)
    
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@tool
def parse_structured_content(structured_content_json: str) -> str:
    """
    解析前端传来的结构化内容
    
    前端会发送一个JSON，描述用户输入的每个内容块及其对应的样式类型。
    这个工具解析并验证该结构，返回可供代码生成使用的规范格式。
    
    Args:
        structured_content_json: 结构化内容的JSON字符串
            格式示例：
            {
                "elements": [
                    {"type": "heading1", "style_name": "Heading 1", "text": "文章标题"},
                    {"type": "heading2", "style_name": "Heading 2", "text": "第一章"},
                    {"type": "body", "style_name": "Normal", "text": "正文内容..."},
                    {"type": "image", "placeholder": "图1说明"}
                ]
            }
        
    Returns:
        解析后的结构化内容描述
    """
    try:
        content = json.loads(structured_content_json)
        elements = content.get("elements", [])
        
        if not elements:
            return "错误：内容为空，请提供至少一个内容元素"
        
        # 验证和规范化
        validated_elements = []
        for i, elem in enumerate(elements):
            elem_type = elem.get("type", "body")
            style_name = elem.get("style_name", "Normal")
            text = elem.get("text", "")
            
            validated_elements.append({
                "index": i,
                "type": elem_type,
                "style_name": style_name,
                "text": text[:100] + "..." if len(text) > 100 else text,
                "text_length": len(text)
            })
        
        result = {
            "total_elements": len(validated_elements),
            "elements": validated_elements,
            "summary": f"共 {len(validated_elements)} 个内容块"
        }
        
        # 统计各类型数量
        type_counts = {}
        for elem in validated_elements:
            t = elem["type"]
            type_counts[t] = type_counts.get(t, 0) + 1
        
        result["type_counts"] = type_counts
        
        return json.dumps(result, ensure_ascii=False, indent=2)
    
    except json.JSONDecodeError as e:
        return f"JSON解析错误：{str(e)}"
    except Exception as e:
        return f"解析结构化内容失败：{str(e)}"


@tool
def analyze_content_structure(file_path: str) -> str:
    """
    智能分析内容文档的结构，识别标题、正文、图片、表格
    
    这个工具会自动识别：
    - 标题（通常是前几段、较短、可能加粗或字号较大）
    - 副标题
    - 正文段落
    - 图片位置和数量
    - 表格位置和数量
    
    Args:
        file_path: 内容文件路径
        
    Returns:
        内容结构的简洁摘要
    """
    try:
        doc = Document(file_path)
        
        # 分析所有段落
        titles = []
        body_paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            
            text = para.text.strip()
            
            # 识别标题的启发式规则：
            # 1. 前3段
            # 2. 字数较少（<50字）
            # 3. 加粗或字号较大
            is_likely_title = False
            
            if i < 3:  # 前3段更可能是标题
                if len(text) < 50:  # 短句
                    is_likely_title = True
                elif para.runs and para.runs[0].bold:  # 加粗
                    is_likely_title = True
                elif para.runs and para.runs[0].font.size and para.runs[0].font.size.pt > 14:  # 大字号
                    is_likely_title = True
            
            if is_likely_title:
                font_size = para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else "未知"
                is_bold = para.runs[0].bold if para.runs else False
                titles.append({
                    "position": i,
                    "text": text[:50],
                    "font_size": font_size,
                    "bold": is_bold
                })
            else:
                body_paragraphs.append(i)
        
        # 统计图片
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
        
        # 统计表格
        table_count = len(doc.tables)
        
        # 生成摘要
        summary_parts = [f"[OK] 内容分析完成（共 {len(doc.paragraphs)} 个段落）：\n"]
        
        # 标题信息
        if titles:
            summary_parts.append(f"[标题] 识别到 {len(titles)} 个标题：")
            for i, title in enumerate(titles[:3]):  # 只显示前3个
                font_info = f"{title['font_size']}pt" if title['font_size'] != "未知" else "默认大小"
                bold_info = "粗体" if title['bold'] else "普通"
                summary_parts.append(f"  {i+1}. \"{title['text']}...\" ({font_info}, {bold_info})")
        else:
            summary_parts.append("[提示] 未识别到明显的标题（可能全文都是正文）")
        
        # 正文信息
        summary_parts.append(f"\n[正文] 正文段落：{len(body_paragraphs)} 段")
        
        # 图片和表格
        if image_count > 0:
            summary_parts.append(f"🖼️  包含 {image_count} 张图片")
        if table_count > 0:
            summary_parts.append(f"[表格] 包含 {table_count} 个表格")
        
        return "\n".join(summary_parts)
    
    except Exception as e:
        return f"分析内容失败：{str(e)}"

@tool
def summarize_document_structure(structure_json: str, max_items: int = 10) -> str:
    """
    将文档结构 JSON 转换为简洁的摘要
    
    ⚠️ 废弃：请使用 analyze_template_structure 直接分析文件
    
    这个工具帮助你生成一个简短的文档结构摘要，避免 JSON 解析错误。
    
    Args:
        structure_json: extract_document_structure 返回的 JSON 字符串
        max_items: 要分析的最大元素数量（默认10个）
        
    Returns:
        简洁的文档结构摘要
    """
    try:
        # 解析 JSON
        if isinstance(structure_json, str):
            items = json.loads(structure_json)
        else:
            items = structure_json
        
        # 只分析前 N 个元素
        items = items[:max_items] if isinstance(items, list) else []
        
        summary_parts = []
        
        # 统计样式
        styles = {}
        for item in items:
            style = item.get('style', 'Unknown')
            if style not in styles:
                styles[style] = []
            styles[style].append(item)
        
        # 生成摘要
        summary_parts.append(f"文档包含 {len(items)} 个主要元素：")
        
        for style, items_list in list(styles.items())[:3]:  # 只显示前3种样式
            sample = items_list[0]
            summary_parts.append(
                f"- {style} 样式：{sample.get('font_size', '未知')}pt, "
                f"{'粗体' if sample.get('bold') else '普通'}, "
                f"对齐方式: {sample.get('alignment', '未知')}"
            )
            if sample.get('text'):
                summary_parts.append(f"  示例文本：{sample['text'][:30]}...")
        
        return "\n".join(summary_parts)
    
    except Exception as e:
        return f"生成摘要失败：{str(e)}"

@tool
def read_document(file_path: str) -> str:
    """
    读取 Word 文档内容
    
    Args:
        file_path: Word 文档的完整路径
        
    Returns:
        文档的文本内容
    """
    try:
        doc = Document(file_path)
        
        # 提取所有段落
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # 提取表格
        tables_content = []
        for table in doc.tables:
            table_text = "\n".join([
                "\t".join([cell.text for cell in row.cells])
                for row in table.rows
            ])
            tables_content.append(f"[表格]\n{table_text}")
        
        content = "\n\n".join(paragraphs + tables_content)
        
        return f"成功读取文档，内容如下：\n\n{content[:2000]}..."  # 限制长度避免 token 溢出
    
    except Exception as e:
        return f"读取文档失败：{str(e)}"

@tool
def write_document(file_path: str, content: str, title: Optional[str] = None) -> str:
    """
    创建新的 Word 文档
    
    Args:
        file_path: 保存路径
        content: 文档内容（支持段落分隔符 \\n\\n）
        title: 可选的文档标题
        
    Returns:
        操作结果
    """
    try:
        doc = Document()
        
        # 添加标题
        if title:
            doc.add_heading(title, level=0)
        
        # 添加内容（按段落分割）
        paragraphs = content.split("\n\n")
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # 保存
        output_path = Path(file_path).with_suffix('.docx')
        doc.save(str(output_path))
        
        return f"成功创建文档：{output_path}"
    
    except Exception as e:
        return f"创建文档失败：{str(e)}"

@tool
def modify_document(
    file_path: str,
    operation: str,
    target: Optional[str] = None,
    replacement: Optional[str] = None
) -> str:
    """
    修改现有 Word 文档
    
    Args:
        file_path: 文档路径
        operation: 操作类型（replace_text, add_paragraph, change_style）
        target: 目标内容（如要替换的文本）
        replacement: 替换内容
        
    Returns:
        操作结果
    """
    try:
        doc = Document(file_path)
        
        if operation == "replace_text" and target and replacement:
            # 替换文本
            count = 0
            for para in doc.paragraphs:
                if target in para.text:
                    para.text = para.text.replace(target, replacement)
                    count += 1
            
            # 保存修改
            output_path = Path(file_path).with_stem(
                Path(file_path).stem + "_modified"
            )
            doc.save(str(output_path))
            
            return f"成功替换 {count} 处文本，保存至：{output_path}"
        
        elif operation == "add_paragraph":
            # 添加新段落
            doc.add_paragraph(replacement or "")
            output_path = Path(file_path).with_stem(
                Path(file_path).stem + "_modified"
            )
            doc.save(str(output_path))
            return f"成功添加段落，保存至：{output_path}"
        
        else:
            return f"不支持的操作：{operation}"
    
    except Exception as e:
        return f"修改文档失败：{str(e)}"

# ============ Markdown 处理工具 ============

@tool
def read_markdown(file_path: str) -> str:
    """
    读取 Markdown 文件内容和结构
    
    Args:
        file_path: Markdown 文件路径
        
    Returns:
        文档内容和结构信息
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取标题结构
        headers = re.findall(r'^(#{1,6})\s+(.+)$', content, re.MULTILINE)
        structure = [{"level": len(h[0]), "text": h[1]} for h in headers]
        
        # 提取图片引用
        images = re.findall(r'!\[([^\]]*)\]\(([^\)]+)\)', content)
        
        result = {
            "content": content[:2000],  # 限制长度
            "structure": structure,
            "images": [{"alt": img[0], "path": img[1]} for img in images],
            "lines": len(content.split('\n'))
        }
        
        return f"成功读取 Markdown 文件：\n{json.dumps(result, ensure_ascii=False, indent=2)}"
    
    except Exception as e:
        return f"读取 Markdown 失败：{str(e)}"

@tool
def write_markdown(
    file_path: str,
    content: str,
    title: Optional[str] = None,
    metadata: Optional[str] = None
) -> str:
    """
    创建 Markdown 文件
    
    Args:
        file_path: 保存路径
        content: Markdown 内容
        title: 文档标题（会作为一级标题）
        metadata: YAML 格式的元数据（可选）
        
    Returns:
        操作结果
    """
    try:
        md_content = ""
        
        # 添加元数据
        if metadata:
            md_content += f"---\n{metadata}\n---\n\n"
        
        # 添加标题
        if title:
            md_content += f"# {title}\n\n"
        
        # 添加内容
        md_content += content
        
        # 保存文件
        output_path = Path(file_path).with_suffix('.md')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return f"成功创建 Markdown 文件：{output_path}"
    
    except Exception as e:
        return f"创建 Markdown 失败：{str(e)}"

# ============ LaTeX 处理工具 ============

@tool
def read_latex(file_path: str) -> str:
    """
    读取 LaTeX 文件内容和结构
    
    Args:
        file_path: LaTeX 文件路径
        
    Returns:
        文档内容和结构信息
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取文档类
        doc_class = re.search(r'\\documentclass\[?([^\]]*)\]?\{([^\}]+)\}', content)
        
        # 提取章节结构
        sections = re.findall(
            r'\\(chapter|section|subsection|subsubsection)\{([^\}]+)\}',
            content
        )
        
        # 提取图片
        figures = re.findall(r'\\includegraphics\[?([^\]]*)\]?\{([^\}]+)\}', content)
        
        result = {
            "document_class": doc_class.group(2) if doc_class else "unknown",
            "sections": [{"type": s[0], "title": s[1]} for s in sections],
            "figures": [{"options": f[0], "path": f[1]} for f in figures],
            "content_preview": content[:1000]
        }
        
        return f"成功读取 LaTeX 文件：\n{json.dumps(result, ensure_ascii=False, indent=2)}"
    
    except Exception as e:
        return f"读取 LaTeX 失败：{str(e)}"

@tool
def write_latex(
    file_path: str,
    content: str,
    document_class: str = "article",
    packages: Optional[List[str]] = None,
    title: Optional[str] = None,
    author: Optional[str] = None
) -> str:
    """
    创建 LaTeX 文件
    
    Args:
        file_path: 保存路径
        content: LaTeX 内容（document 环境内的内容）
        document_class: 文档类（article, report, book 等）
        packages: 需要引入的包列表
        title: 文档标题
        author: 作者
        
    Returns:
        操作结果
    """
    try:
        # 构建 LaTeX 文档
        latex_content = f"\\documentclass{{{document_class}}}\n\n"
        
        # 添加常用包
        default_packages = ['graphicx', 'amsmath', 'geometry', 'xeCJK']
        all_packages = default_packages + (packages or [])
        
        for pkg in all_packages:
            latex_content += f"\\usepackage{{{pkg}}}\n"
        
        latex_content += "\n"
        
        # 添加标题和作者
        if title:
            latex_content += f"\\title{{{title}}}\n"
        if author:
            latex_content += f"\\author{{{author}}}\n"
        
        latex_content += f"\\date{{\\today}}\n\n"
        latex_content += "\\begin{document}\n\n"
        
        if title:
            latex_content += "\\maketitle\n\n"
        
        # 添加正文内容
        latex_content += content
        
        latex_content += "\n\\end{document}"
        
        # 保存文件
        output_path = Path(file_path).with_suffix('.tex')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(latex_content)
        
        return f"成功创建 LaTeX 文件：{output_path}"
    
    except Exception as e:
        return f"创建 LaTeX 失败：{str(e)}"

# ============ 图片处理工具 ============

@tool
def extract_images_from_document(file_path: str, output_dir: str = "extracted_images") -> str:
    """
    从 Word 文档中提取所有图片
    
    Args:
        file_path: Word 文档路径
        output_dir: 图片保存目录
        
    Returns:
        提取的图片信息
    """
    try:
        doc = Document(file_path)
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        images_info = []
        
        # 遍历文档中的所有关系（包含图片）
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                # 提取图片
                image_data = rel.target_part.blob
                
                # 确定文件扩展名
                ext = rel.target_ref.split('.')[-1]
                image_name = f"image_{len(images_info) + 1}.{ext}"
                image_path = output_path / image_name
                
                # 保存图片
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                
                # 获取图片尺寸
                img = Image.open(BytesIO(image_data))
                
                images_info.append({
                    "name": image_name,
                    "path": str(image_path),
                    "size": img.size,
                    "format": img.format
                })
        
        return f"成功提取 {len(images_info)} 张图片：\n{json.dumps(images_info, ensure_ascii=False, indent=2)}"
    
    except Exception as e:
        return f"提取图片失败：{str(e)}"

# ============ 文档结构分析工具 ============

@tool
def extract_document_structure(file_path: str) -> str:
    """
    提取文档的结构信息（标题、段落、样式等）
    
    Args:
        file_path: 文档路径（支持 .doc, .docx, .md, .tex）
        
    Returns:
        文档结构的 JSON 表示
    """
    try:
        file_ext = Path(file_path).suffix.lower()
        actual_file_path = file_path
        
        # 如果是 .doc 格式，先尝试转换为 .docx
        if file_ext == '.doc':
            docx_path = Path(file_path).with_suffix('.docx')
            result = convert_doc_to_docx(file_path, str(docx_path))
            if "成功" in result:
                actual_file_path = str(docx_path)
                file_ext = '.docx'
            else:
                return f"无法处理 .doc 格式文件。请先将文件转换为 .docx 格式。\n提示：可以使用 Microsoft Word 打开文件后另存为 .docx 格式，或使用 convert_format 工具转换。\n转换尝试结果：{result}"
        
        if file_ext == '.docx':
            doc = Document(actual_file_path)
            structure = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # 分析样式
                    style_info = {
                        "text": para.text[:100],
                        "style": para.style.name,
                        "is_heading": para.style.name.startswith('Heading'),
                        "alignment": str(para.alignment),
                        "font_size": para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else None,
                        "bold": para.runs[0].bold if para.runs else False,
                        "italic": para.runs[0].italic if para.runs else False
                    }
                    structure.append(style_info)
            
            return f"Word 文档结构（共 {len(structure)} 个元素）：\n{json.dumps(structure[:10], ensure_ascii=False, indent=2)}"
        
        elif file_ext == '.md':
            return read_markdown(actual_file_path)
        
        elif file_ext == '.tex':
            return read_latex(actual_file_path)
        
        else:
            return f"不支持的文件格式：{file_ext}"
    
    except Exception as e:
        return f"提取结构失败：{str(e)}"

# ============ 模板应用工具 ============

@tool
def apply_template_structure(
    template_structure: str,
    content_text: str,
    output_path: str,
    output_format: str = "docx"
) -> str:
    """
    将内容应用到模板结构中
    
    Args:
        template_structure: 模板的结构信息（JSON 字符串）
        content_text: 要填充的内容
        output_path: 输出文件路径
        output_format: 输出格式（docx/md/tex）
        
    Returns:
        操作结果
    """
    try:
        structure = json.loads(template_structure)
        
        if output_format == "docx":
            doc = Document()
            
            # 按照模板结构创建文档
            content_paragraphs = content_text.split('\n\n')
            
            for i, item in enumerate(structure):
                if i < len(content_paragraphs):
                    para = doc.add_paragraph(content_paragraphs[i])
                    
                    # 应用样式
                    if item.get('is_heading'):
                        para.style = item.get('style', 'Normal')
                    
                    if item.get('bold') and para.runs:
                        para.runs[0].bold = True
                    
                    if item.get('italic') and para.runs:
                        para.runs[0].italic = True
            
            doc.save(output_path)
            return f"成功应用模板并保存到：{output_path}"
        
        elif output_format == "md":
            # 构建 Markdown
            md_lines = []
            content_paragraphs = content_text.split('\n\n')
            
            for i, item in enumerate(structure):
                if i < len(content_paragraphs):
                    text = content_paragraphs[i]
                    
                    # 根据原结构添加格式
                    if item.get('level'):
                        prefix = '#' * item['level']
                        md_lines.append(f"{prefix} {text}")
                    else:
                        md_lines.append(text)
                    
                    md_lines.append("")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_lines))
            
            return f"成功生成 Markdown：{output_path}"
        
        else:
            return f"暂不支持输出格式：{output_format}"
    
    except Exception as e:
        return f"应用模板失败：{str(e)}"

# ============ 格式转换工具 ============

def convert_doc_to_docx(doc_path: str, docx_path: str) -> str:
    """
    尝试将 .doc 格式转换为 .docx 格式（辅助函数，不是工具）
    
    Args:
        doc_path: 源 .doc 文件路径
        docx_path: 目标 .docx 文件路径
        
    Returns:
        转换结果消息
    """
    import subprocess
    import platform
    
    try:
        system = platform.system()
        
        # 方法 1: 尝试使用 LibreOffice (跨平台)
        if system == "Windows":
            # Windows 上尝试 LibreOffice
            libreoffice_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
            ]
            for lo_path in libreoffice_paths:
                if Path(lo_path).exists():
                    output_dir = Path(doc_path).parent
                    result = subprocess.run(
                        [lo_path, "--headless", "--convert-to", "docx", "--outdir", str(output_dir), doc_path],
                        capture_output=True,
                        timeout=30
                    )
                    if result.returncode == 0:
                        # LibreOffice 会生成同名的 .docx 文件
                        generated_docx = Path(doc_path).with_suffix('.docx')
                        if generated_docx.exists() and str(generated_docx) != docx_path:
                            import shutil
                            shutil.move(str(generated_docx), docx_path)
                        return f"成功使用 LibreOffice 转换 .doc 为 .docx：{docx_path}"
        
        # 如果 LibreOffice 不可用，返回提示信息
        return f"未找到可用的转换工具。请手动将 .doc 文件转换为 .docx 格式。\n提示：安装 LibreOffice 后可自动转换。"
    
    except subprocess.TimeoutExpired:
        return "转换超时。文件可能过大或转换工具响应缓慢。"
    except Exception as e:
        return f"转换过程出错：{str(e)}"

@tool
def convert_format(
    input_path: str,
    output_format: str,
    preserve_images: bool = True
) -> str:
    """
    转换文档格式
    
    Args:
        input_path: 输入文件路径
        output_format: 目标格式（word/markdown/latex）
        preserve_images: 是否保留图片
        
    Returns:
        转换结果和输出文件路径
    """
    try:
        input_ext = Path(input_path).suffix.lower()
        output_path = Path(input_path).with_suffix(f'.{output_format}')
        
        # 如果输入是 .doc，先转换为 .docx
        actual_input_path = input_path
        if input_ext == '.doc':
            temp_docx = Path(input_path).with_suffix('.docx')
            conversion_result = convert_doc_to_docx(input_path, str(temp_docx))
            if "成功" in conversion_result:
                actual_input_path = str(temp_docx)
                input_ext = '.docx'
            else:
                return f"无法转换 .doc 文件：{conversion_result}"
        
        # 读取原文档
        if input_ext == '.docx':
            doc = Document(actual_input_path)
            text_content = '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        elif input_ext == '.md':
            with open(actual_input_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
        else:
            return f"暂不支持从 {input_ext} 格式转换"
        
        # 转换为目标格式
        if output_format == 'markdown' or output_format == 'md':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            return f"成功转换为 Markdown：{output_path}"
        
        elif output_format == 'word' or output_format == 'docx':
            new_doc = Document()
            for para in text_content.split('\n\n'):
                if para.strip():
                    new_doc.add_paragraph(para.strip())
            new_doc.save(str(output_path))
            return f"成功转换为 Word：{output_path}"
        
        else:
            return f"不支持的输出格式：{output_format}"
    
    except Exception as e:
        return f"格式转换失败：{str(e)}"

def get_document_tools():
    """返回所有工具的列表"""
    return [
        analyze_template_structure,  # 模板分析（增强版）
        get_template_styles,  # 获取模板样式列表（新增）
        parse_structured_content,  # 解析结构化内容（新增）
        analyze_content_structure,  # 内容分析
        summarize_document_structure,  # 保留但标记为废弃
        read_document,
        write_document,
        modify_document,
        read_markdown,
        write_markdown,
        read_latex,
        write_latex,
        extract_images_from_document,
        extract_document_structure,
        apply_template_structure,
        convert_format
    ]

