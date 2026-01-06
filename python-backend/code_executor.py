"""
代码生成和执行工具
让 LLM 生成 Python 代码来处理文档，而不是直接操作
"""
from langchain.tools import tool
import subprocess
import sys
from pathlib import Path
import tempfile
import os
import re

@tool
def generate_document_processing_code(
    template_summary: str,
    content_summary: str,
    output_path: str
) -> str:
    """
    生成文档处理代码的指引工具
    
    ⚠️ 重要：参数必须是简短的摘要文字，不要传递完整的 JSON 结构！
    
    这个工具会返回一个提示，指导 AI 如何生成 Python 代码来处理文档。
    
    Args:
        template_summary: 模板文件的摘要信息（3-5 句话描述样式、格式、结构）
        content_summary: 内容文件的摘要信息（标题、主要内容概述）
        output_path: 输出文件的完整路径（例如：uploads/output_xxx_20240106.docx）
        
    Returns:
        代码生成指引
        
    示例调用：
        template_summary="模板是大连理工大学论文格式，包含标题（24pt粗体居中）、副标题（22pt）、学院/专业等信息字段（15pt）"
        content_summary="内容是小说《突然放弃的习惯》，包含标题、副标题和正文段落"
        output_path="uploads/output_突然放弃的习惯_20240106_203945.docx"
    """
    guide = f"""
✅ 现在请生成完整的 Python 代码来处理文档。

【模板摘要】
{template_summary}

【内容摘要】
{content_summary}

【输出路径】
{output_path}

【代码要求】
1. 使用 python-docx 库操作 Word 文档
2. 代码必须完整可执行，包含所有 import 语句
3. 从模板文件加载：Document("模板文件路径")
4. 清空模板中的所有段落（保留样式）
5. 将内容按照模板样式填入新段落
6. 保存到输出路径：doc.save("OUTPUT_PATH_PLACEHOLDER")
7. 不要使用硬编码的路径，使用占位符

【代码框架示例】
```python
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. 加载模板
doc = Document("模板文件路径")

# 2. 读取内容文档
content_doc = Document("内容文件路径")
content_paragraphs = [p.text for p in content_doc.paragraphs if p.text.strip()]

# 3. 获取模板的样式（从第一个段落）
template_styles = {{
    'title': doc.paragraphs[0].style if len(doc.paragraphs) > 0 else None,
    'body': doc.paragraphs[-1].style if len(doc.paragraphs) > 0 else None
}}

# 4. 清空模板内容（保留第一页的格式信息）
# 从后往前删除，保留前几个格式段落
for i in range(len(doc.paragraphs) - 1, 10, -1):  # 保留前10个段落作为格式参考
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)

# 5. 添加内容（使用模板样式）
for i, text in enumerate(content_paragraphs):
    if i == 0:  # 标题
        p = doc.add_paragraph(text)
        p.style = template_styles['title']
    else:  # 正文
        p = doc.add_paragraph(text)
        p.style = template_styles['body']

# 6. 保存
doc.save("OUTPUT_PATH_PLACEHOLDER")
print("文档已生成")
```

请直接输出完整的 Python 代码，不要有额外的解释。代码要用 ```python 包裹。

⚠️⚠️⚠️ 生成代码后的必须执行步骤（不要跳过！）：

生成代码后，你**必须立即**调用 execute_generated_code 工具来执行这段代码！

参数如下：
- code: 你刚生成的完整 Python 代码（包含 ```python 和 ``` 标记）
- template_path: "模板文件路径"（从用户指令中提取，应该是 uploads\\template_xxx.docx）
- content_path: "内容文件路径"（从用户指令中提取，应该是 uploads\\content_xxx.docx）
- output_path: "{output_path}"（就是上面的输出路径）

如果不执行代码，文档不会被生成！用户会看到"未找到输出文件"错误！

⚠️ 下一步必须：调用 execute_generated_code 执行代码！
"""
    return guide

@tool
def execute_generated_code(
    code: str,
    template_path: str,
    content_path: str,
    output_path: str
) -> str:
    """
    在安全的环境中执行 LLM 生成的代码
    
    Args:
        code: LLM 生成的 Python 代码
        template_path: 模板文件路径
        content_path: 内容文件路径
        output_path: 输出文件路径
        
    Returns:
        执行结果
    """
    try:
        # 提取代码块
        code_match = re.search(r'```python\s*(.*?)\s*```', code, re.DOTALL)
        if code_match:
            clean_code = code_match.group(1)
        else:
            clean_code = code
        
        # 输出调试信息
        print("\n" + "="*60)
        print("【代码执行调试信息】")
        print(f"模板路径: {template_path}")
        print(f"内容路径: {content_path}")
        print(f"输出路径: {output_path}")
        print("="*60 + "\n")
        
        # 替换代码中的路径占位符（多种变体）
        clean_code = clean_code.replace('"模板文件路径"', f'r"{template_path}"')
        clean_code = clean_code.replace('"内容文件路径"', f'r"{content_path}"')
        
        # 替换输出路径的各种形式
        clean_code = clean_code.replace('"OUTPUT_PATH_PLACEHOLDER"', f'r"{output_path}"')
        clean_code = clean_code.replace("'OUTPUT_PATH_PLACEHOLDER'", f'r"{output_path}"')
        clean_code = clean_code.replace('"output.docx"', f'r"{output_path}"')
        clean_code = clean_code.replace("'output.docx'", f'r"{output_path}"')
        clean_code = clean_code.replace('"uploads/output.docx"', f'r"{output_path}"')
        clean_code = clean_code.replace("'uploads/output.docx'", f'r"{output_path}"')
        
        # 替换具体的路径引用（如果代码中直接使用了路径）
        clean_code = re.sub(
            r'output_path\s*=\s*["\'][^"\']*?\.docx["\']',
            f'output_path = r"{output_path}"',
            clean_code
        )
        
        # 输出替换后的代码（用于调试）
        print("【替换路径后的代码】")
        print(clean_code)
        print("="*60 + "\n")
        
        # 确保输出目录存在
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # 获取项目根目录（python-backend的父目录）
        project_root = Path(__file__).parent.parent.resolve()
        
        # 创建临时文件保存代码
        with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False, encoding='utf-8') as f:
            temp_script = f.name
            f.write(clean_code)
        
        try:
            # 执行代码（设置超时防止死循环）
            result = subprocess.run(
                [sys.executable, temp_script],
                capture_output=True,
                text=True,
                timeout=60,  # 60秒超时
                cwd=str(project_root)  # 在项目根目录执行
            )
            
            # 使用绝对路径检查文件是否存在
            output_file = Path(output_path).resolve()
            
            # 检查是否成功
            if result.returncode == 0:
                if output_file.exists():
                    return f"✅ 代码执行成功！文档已生成：{output_file}\n\n执行输出：\n{result.stdout}"
                else:
                    return f"⚠️ 代码执行完成，但未找到输出文件。\n期望路径：{output_file}\n\n输出：\n{result.stdout}\n\n错误：\n{result.stderr}"
            else:
                return f"❌ 代码执行失败：\n\n{result.stderr}\n\n输出：\n{result.stdout}"
        
        finally:
            # 清理临时文件
            if os.path.exists(temp_script):
                os.remove(temp_script)
    
    except subprocess.TimeoutExpired:
        return "❌ 代码执行超时（超过60秒），可能存在死循环"
    
    except Exception as e:
        return f"❌ 执行代码时出错：{str(e)}"

@tool
def validate_generated_code(code: str) -> str:
    """
    验证生成的代码是否安全
    
    Args:
        code: 待验证的 Python 代码
        
    Returns:
        验证结果
    """
    # 危险操作列表
    dangerous_patterns = [
        r'import\s+os\s*\.',  # os 模块的危险操作
        r'os\.(system|popen|remove|rmdir)',
        r'subprocess\.(call|run|Popen)',
        r'eval\s*\(',
        r'exec\s*\(',
        r'__import__',
        r'open\s*\([^)]*[\'"]w',  # 写入非预期文件
        r'shutil\.(rmtree|move)',
    ]
    
    issues = []
    for pattern in dangerous_patterns:
        if re.search(pattern, code):
            issues.append(f"发现潜在危险操作：{pattern}")
    
    # 检查必要的库
    required_imports = ['from docx import', 'import docx']
    has_docx = any(imp in code for imp in required_imports)
    
    if not has_docx:
        issues.append("代码中未导入 python-docx 库")
    
    if issues:
        return "⚠️ 代码验证警告：\n" + "\n".join(f"- {issue}" for issue in issues)
    else:
        return "✅ 代码验证通过，可以安全执行"

@tool
def generate_and_execute_document_code(
    template_path: str,
    content_path: str,
    output_path: str,
    template_summary: str,
    content_summary: str
) -> str:
    """
    ⭐ 一键生成并执行文档处理代码（推荐使用！）
    
    这个工具会自动完成：
    1. 生成 Python 代码来处理文档
    2. 立即执行生成的代码
    3. 返回执行结果
    
    使用这个工具可以避免遗忘执行代码的步骤。
    
    Args:
        template_path: 模板文件的完整路径（从用户指令中获取）
        content_path: 内容文件的完整路径（从用户指令中获取）
        output_path: 输出文件路径（从用户指令中获取）
        template_summary: 模板的简短摘要（3-5句话描述样式和格式）
        content_summary: 内容的简短摘要（标题和主要内容）
        
    Returns:
        执行结果
        
    示例调用：
        template_path="uploads/template_xxx.docx"
        content_path="uploads/content_xxx.docx"
        output_path="uploads/output_xxx_20240106.docx"
        template_summary="模板是大工论文格式，标题24pt粗体居中，副标题22pt，正文15pt"
        content_summary="内容是小说《平滑度过的人生》，包含标题、副标题和正文段落"
    """
    try:
        # 生成固定的代码模板（简化版）
        generated_code = """```python
from docx import Document

# 加载模板
doc = Document("模板文件路径")

# 读取内容
content_doc = Document("内容文件路径")
content_texts = [p.text for p in content_doc.paragraphs if p.text.strip()]

# 获取模板样式
title_style = doc.paragraphs[0].style if len(doc.paragraphs) > 0 else None
body_style = doc.paragraphs[10].style if len(doc.paragraphs) > 10 else None

# 清空模板（保留前10个段落作为格式参考）
for i in range(len(doc.paragraphs) - 1, 10, -1):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)

# 添加内容
for i, text in enumerate(content_texts):
    p = doc.add_paragraph(text)
    p.style = title_style if i == 0 else body_style

# 保存
doc.save("OUTPUT_PATH_PLACEHOLDER")
print("文档已生成")
```"""
        
        # 内联执行代码逻辑（不调用工具）
        try:
            # 提取代码块
            code_match = re.search(r'```python\s*(.*?)\s*```', generated_code, re.DOTALL)
            if code_match:
                clean_code = code_match.group(1)
            else:
                clean_code = generated_code
            
            # 输出调试信息
            print("\n" + "="*60)
            print("【代码执行调试信息】")
            print(f"模板路径: {template_path}")
            print(f"内容路径: {content_path}")
            print(f"输出路径: {output_path}")
            print("="*60 + "\n")
            
            # 转换为绝对路径（关键修复！）
            # 注意：uploads 文件夹在 python-backend 目录下
            backend_root = Path(__file__).parent.resolve()  # python-backend 目录
            abs_template_path = (backend_root / template_path).resolve()
            abs_content_path = (backend_root / content_path).resolve()
            abs_output_path = (backend_root / output_path).resolve()
            
            print("【绝对路径】")
            print(f"模板: {abs_template_path}")
            print(f"内容: {abs_content_path}")
            print(f"输出: {abs_output_path}")
            
            # 检查文件是否存在
            if not abs_template_path.exists():
                return f"❌ 模板文件不存在：{abs_template_path}"
            if not abs_content_path.exists():
                return f"❌ 内容文件不存在：{abs_content_path}"
            
            print("✅ 文件检查通过")
            print("="*60 + "\n")
            
            # 替换路径占位符（使用绝对路径）
            clean_code = clean_code.replace('"模板文件路径"', f'r"{abs_template_path}"')
            clean_code = clean_code.replace('"内容文件路径"', f'r"{abs_content_path}"')
            clean_code = clean_code.replace('"OUTPUT_PATH_PLACEHOLDER"', f'r"{abs_output_path}"')
            clean_code = clean_code.replace("'OUTPUT_PATH_PLACEHOLDER'", f'r"{abs_output_path}"')
            
            # 输出替换后的代码
            print("【替换路径后的代码】")
            print(clean_code)
            print("="*60 + "\n")
            
            # 确保输出目录存在
            abs_output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 创建临时文件保存代码
            with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False, encoding='utf-8') as f:
                temp_script = f.name
                f.write(clean_code)
            
            try:
                # 执行代码（工作目录设置为 python-backend）
                result = subprocess.run(
                    [sys.executable, temp_script],
                    capture_output=True,
                    text=True,
                    timeout=60,
                    cwd=str(backend_root)  # 使用 python-backend 作为工作目录
                )
                
                # 检查结果（使用绝对路径）
                if result.returncode == 0:
                    if abs_output_path.exists():
                        return f"✅ 代码已自动生成并执行成功！文档已生成：{abs_output_path}\n\n执行输出：\n{result.stdout}"
                    else:
                        return f"⚠️ 代码执行完成，但未找到输出文件。\n期望路径：{abs_output_path}\n\n输出：\n{result.stdout}\n\n错误：\n{result.stderr}"
                else:
                    return f"❌ 代码执行失败：\n\n{result.stderr}\n\n输出：\n{result.stdout}"
            
            finally:
                # 清理临时文件
                if os.path.exists(temp_script):
                    os.remove(temp_script)
        
        except subprocess.TimeoutExpired:
            return "❌ 代码执行超时（超过60秒）"
        except Exception as e:
            return f"❌ 执行过程出错：{str(e)}"
        
    except Exception as e:
        return f"❌ 生成并执行代码失败：{str(e)}"

def get_code_execution_tools():
    """返回代码执行相关的工具"""
    return [
        generate_and_execute_document_code,  # 新的组合工具（推荐）
        generate_document_processing_code,
        execute_generated_code,
        validate_generated_code
    ]
