"""
FastAPI 主应用入口
运行在本地，为前端提供 API 服务
"""
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse
import uvicorn
from pathlib import Path
from typing import Optional
import os
import json
import asyncio
from queue import Queue
import time
import threading
import glob
from datetime import datetime

from agent import DocumentAgent

app = FastAPI(title="ConvertAgent API", version="1.0.0")

# 配置 CORS，允许前端访问
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Tauri 会从 tauri://localhost 访问
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 初始化 Agent（延迟初始化，等待 API Key）
agent: Optional[DocumentAgent] = None

# 进度队列，用于实时推送执行步骤
progress_queue: Queue = Queue()

@app.get("/")
async def root():
    """健康检查"""
    return {"status": "ok", "message": "ConvertAgent Backend is running"}

@app.get("/api/providers")
async def get_providers():
    """
    获取支持的 API 提供商列表
    """
    from agent import DocumentAgent
    providers = []
    for provider_id, config in DocumentAgent.PROVIDERS.items():
        providers.append({
            "id": provider_id,
            "name": config["name"],
            "models": config["models"],
            "default_model": config["default_model"]
        })
    return {"providers": providers}

@app.get("/api/progress")
async def progress_stream():
    """
    SSE 端点，实时推送 Agent 执行进度
    """
    async def event_generator():
        while True:
            # 检查队列中是否有新消息
            if not progress_queue.empty():
                progress_data = progress_queue.get()
                # 发送 SSE 格式的数据
                yield f"data: {json.dumps(progress_data, ensure_ascii=False)}\n\n"
                
                # 如果是完成或错误消息，结束流
                if progress_data.get("type") in ["complete", "error"]:
                    break
            else:
                # 发送心跳保持连接
                yield f": heartbeat\n\n"
            
            await asyncio.sleep(0.1)  # 避免CPU占用过高
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"  # 禁用Nginx缓冲
        }
    )

@app.post("/api/init")
async def init_agent(
    api_key: str = Form(...),
    provider: str = Form(default="openai"),
    model: str = Form(default=None)
):
    """
    初始化 Agent
    
    Args:
        api_key: API 密钥
        provider: 提供商 ID (openai/siliconflow/zhipu/moonshot/deepseek)
        model: 模型名称（可选）
    """
    global agent
    try:
        agent = DocumentAgent(
            api_key=api_key,
            provider=provider,
            model_name=model
        )
        return {
            "success": True, 
            "message": f"Agent 初始化成功（{agent.provider_name} - {agent.model_name}）"
        }
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )

@app.post("/api/analyze-template")
async def analyze_template(
    template_file: UploadFile = File(...)
):
    """
    分析模板文件并返回可用样式列表
    
    用户上传模板后，前端调用此接口获取模板中的样式定义，
    然后让用户在结构化内容编辑器中选择每段内容应使用的样式。
    
    Args:
        template_file: 模板文件 (.docx)
        
    Returns:
        模板样式分析结果
    """
    try:
        # 保存上传的模板文件
        upload_dir = Path("uploads")
        upload_dir.mkdir(exist_ok=True)
        
        template_path = upload_dir / f"template_{template_file.filename}"
        with open(template_path, "wb") as f:
            f.write(await template_file.read())
        
        # 使用 tools.py 中的函数分析模板
        from tools import get_template_styles
        
        # 直接调用工具函数（不通过 Agent）
        styles_json = get_template_styles.invoke(str(template_path))
        
        # 解析 JSON 结果
        import json
        styles = json.loads(styles_json)
        
        if "error" in styles:
            return JSONResponse(
                status_code=500,
                content={"success": False, "error": styles["error"]}
            )
        
        return {
            "success": True,
            "template_path": str(template_path),
            "styles": styles
        }
    
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )



def process_in_background(file_path: str, instruction: str):
    """在后台线程执行 Agent 处理"""
    try:
        # 调用 Agent 处理（带进度回调）
        result = agent.process(
            file_path, 
            instruction, 
            progress_callback=lambda msg: progress_queue.put({
                "type": "step",
                "message": msg,
                "timestamp": time.time()
            })
        )
        
        # 查找生成的输出文件（在uploads目录和当前目录）
        output_files = []
        upload_dir = Path("uploads")
        current_dir = Path(".")
        
        for search_dir in [upload_dir, current_dir]:
            for ext in ['.docx', '.md', '.tex', '.pdf']:
                files = list(search_dir.glob(f'*output*{ext}'))
                output_files.extend(files)
        
        # 移动当前目录下的输出文件到uploads目录
        for f in output_files[:]:
            if f.parent != upload_dir:
                import shutil
                new_path = upload_dir / f.name
                shutil.move(str(f), str(new_path))
                output_files.remove(f)
                output_files.append(new_path)
        
        # 构建结果消息
        if output_files:
            file_links = []
            for f in output_files:
                file_links.append(f"[FILE] [{f.name}](/api/download/{f.name})")
            result_msg = "\n".join(file_links)
        else:
            result_msg = result.get('output', '处理完成')
            # 移除代码块，只保留摘要
            if '```' in result_msg:
                result_msg = result_msg.split('```')[0].strip() or "处理完成，但未找到输出文件"
        
        # 推送完成消息
        progress_queue.put({
            "type": "complete",
            "message": f"[OK] 处理完成！\n\n{result_msg}",
            "timestamp": time.time()
        })
        
    except Exception as e:
        progress_queue.put({
            "type": "error",
            "message": f"[ERROR] 处理失败: {str(e)}",
            "timestamp": time.time()
        })

@app.post("/api/process")
async def process_document(
    file: UploadFile = File(...),
    instruction: str = Form(...)
):
    """
    处理文档
    接收文件 + 用户指令，返回处理结果
    """
    if agent is None:
        return JSONResponse(
            status_code=400,
            content={"success": False, "error": "请先初始化 Agent"}
        )
    
    try:
        # 清空进度队列
        while not progress_queue.empty():
            progress_queue.get()
        
        # 推送开始消息
        progress_queue.put({
            "type": "start",
            "message": "[START] 开始处理文档...",
            "timestamp": time.time()
        })
        
        # 保存上传的文件
        upload_dir = Path("uploads")
        upload_dir.mkdir(exist_ok=True)
        
        file_path = upload_dir / file.filename
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        progress_queue.put({
            "type": "step",
            "message": f"[SAVE] 文件已保存: {file.filename}",
            "timestamp": time.time()
        })
        
        # 在后台线程执行 Agent（这样才能实时推送进度）
        thread = threading.Thread(
            target=process_in_background,
            args=(str(file_path), instruction)
        )
        thread.daemon = True
        thread.start()
        
        return {
            "success": True,
            "message": "处理已开始，请查看进度面板"
        }
    
    except Exception as e:
        progress_queue.put({
            "type": "error",
            "message": f"[ERROR] 处理失败: {str(e)}",
            "timestamp": time.time()
        })
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )

@app.post("/api/process-with-template")
async def process_with_template(
    template_file: UploadFile = File(...),
    content_file: UploadFile = File(...),
    output_format: str = Form(...),
    additional_instruction: str = Form(default="")
):
    """
    使用模板处理文档
    接收模板文件 + 内容文件，生成指定格式的输出
    
    Args:
        template_file: 模板文件（word/markdown/latex）
        content_file: 内容文件（包含文本和图片）
        output_format: 输出格式（word/markdown/latex）
        additional_instruction: 额外的处理指令
    """
    if agent is None:
        return JSONResponse(
            status_code=400,
            content={"success": False, "error": "请先初始化 Agent"}
        )
    
    try:
        # 清空进度队列
        while not progress_queue.empty():
            progress_queue.get()
        
        # 推送开始消息
        progress_queue.put({
            "type": "start",
            "message": "开始处理模板文档...",
            "timestamp": time.time()
        })
        
        # 保存上传的文件
        upload_dir = Path("uploads")
        upload_dir.mkdir(exist_ok=True)
        
        template_path = upload_dir / f"template_{template_file.filename}"
        content_path = upload_dir / f"content_{content_file.filename}"
        
        progress_queue.put({
            "type": "step",
            "message": f"保存模板文件: {template_file.filename}",
            "timestamp": time.time()
        })
        
        with open(template_path, "wb") as f:
            f.write(await template_file.read())
        
        progress_queue.put({
            "type": "step",
            "message": f"保存内容文件: {content_file.filename}",
            "timestamp": time.time()
        })
        
        with open(content_path, "wb") as f:
            f.write(await content_file.read())
        
        # 生成带时间戳的输出文件名（移到外层）
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        content_name = Path(content_path).stem
        output_filename = f"output_{content_name}_{timestamp}.docx"
        output_file_path = str(Path("uploads") / output_filename)
        
        # 构建指令（强调代码生成和执行）
        instruction = f"""
任务：根据模板生成文档（通过编写并执行 Python 代码实现）

【文件信息】
- 模板文件路径：{template_path}
- 内容文件路径：{content_path}
- 输出文件路径：{output_file_path}
- 输出格式：{output_format}

【完整工作流程】
1. 分析模板结构
   - 使用 analyze_template_structure 分析模板文件
   - 获取样式名、字体大小、对齐方式等信息

2. 读取内容文件
   - 使用 read_document 读取内容文件
   - 识别标题、段落结构

3. 生成 Python 代码
   - 调用 generate_document_processing_code 获取代码框架
   - 编写完整的 Python 代码来处理文档
   - 代码要求：
     * 使用 python-docx 库
     * 从模板文件加载：Document("模板文件路径")
     * 清空模板内容但保留样式
     * 按照模板样式填充新内容
     * 保存到输出路径：doc.save("OUTPUT_PATH_PLACEHOLDER")

4. ⚠️ 执行代码（必须完成！）
   - 调用 execute_generated_code 执行你生成的代码
   - 参数：
     * code: 你生成的完整 Python 代码（包含 ```python 标记）
     * template_path: {template_path}
     * content_path: {content_path}
     * output_path: {output_file_path}
   - 等待执行结果并报告给用户

【额外要求】
{additional_instruction if additional_instruction else "按照模板原有格式处理即可"}

⚠️ 重要：
1. 你的任务是生成 Python 代码并执行它！
2. 只有执行代码后才能生成文档文件！
3. 不要生成代码后就结束，必须调用 execute_generated_code！
"""
        
        # 定义后台处理函数
        def process_template_in_background():
            try:
                progress_queue.put({
                    "type": "step",
                    "message": "[START] 启动 Agent 分析模板...",
                    "timestamp": time.time()
                })
                
                progress_queue.put({
                    "type": "step",
                    "message": f"[INFO] 输出文件：{output_filename}",
                    "timestamp": time.time()
                })
                
                # 调用 Agent 处理（带进度回调）
                result = agent.process_with_template(
                    str(template_path), 
                    str(content_path), 
                    output_format,
                    instruction,  # 使用原始指令，已包含输出路径
                    output_path=output_file_path,  # 传递输出路径
                    progress_callback=lambda msg: progress_queue.put({
                        "type": "step",
                        "message": msg,
                        "timestamp": time.time()
                    })
                )
                
                # 查找生成的输出文件（优先查找指定的输出文件）
                output_files = []
                uploads_dir = Path("uploads")
                
                # 首先检查指定的输出文件
                expected_output = Path(output_file_path)
                if expected_output.exists():
                    output_files.append(expected_output)
                
                # 如果没有找到，再搜索其他可能的输出文件
                # 只查找当前生成的文件（基于时间戳）
                if not output_files:
                    for ext in ['.docx', '.md', '.tex', '.pdf']:
                        # 只查找包含时间戳的文件（最近5分钟内的）
                        import time as time_module
                        current_time = time_module.time()
                        for f in uploads_dir.glob(f'output_*{ext}'):
                            if current_time - f.stat().st_mtime < 300:  # 5分钟内
                                output_files.append(f)
                
                # 构建结果消息
                if output_files:
                    file_links = []
                    for f in output_files:
                        file_links.append(f"📄 [{f.name}](/api/download/{f.name})")
                    result_msg = "\n".join(file_links)
                else:
                    result_msg = result.get('output', '处理完成')
                    # 移除代码块，只保留摘要
                    if '```' in result_msg:
                        result_msg = result_msg.split('```')[0].strip() or "处理完成，但未找到输出文件"
                
                # 推送完成消息
                progress_queue.put({
                    "type": "complete",
                    "message": f"[OK] 模板处理完成！\n\n{result_msg}",
                    "timestamp": time.time()
                })
                
            except Exception as e:
                progress_queue.put({
                    "type": "error",
                    "message": f"[ERROR] 处理失败: {str(e)}",
                    "timestamp": time.time()
                })
        
        # 在后台线程执行
        thread = threading.Thread(target=process_template_in_background)
        thread.daemon = True
        thread.start()
        
        return {
            "success": True,
            "message": "模板处理已开始，请查看进度面板"
        }
    
    except Exception as e:
        progress_queue.put({
            "type": "error",
            "message": f"[ERROR] 处理失败: {str(e)}",
            "timestamp": time.time()
        })
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )

@app.post("/api/process-structured")
async def process_structured(
    template_path: str = Form(...),
    structured_content: str = Form(...),
    output_format: str = Form(default="word")
):
    """
    使用结构化内容生成文档（新流程）
    
    用户通过前端编辑器创建结构化内容（每段内容标记好对应的样式），
    后端直接生成Python代码执行，不需要Agent多步推理。
    
    Args:
        template_path: 模板文件路径（来自analyze-template返回）
        structured_content: 结构化内容JSON
        output_format: 输出格式
        
    JSON格式示例：
        {
            "elements": [
                {"style_name": "Heading 1", "text": "文章标题"},
                {"style_name": "Heading 2", "text": "第一章"},
                {"style_name": "Normal", "text": "正文内容..."}
            ]
        }
    """
    try:
        # 清空进度队列
        while not progress_queue.empty():
            progress_queue.get()
        
        progress_queue.put({
            "type": "start",
            "message": "[START] 开始生成文档...",
            "timestamp": time.time()
        })
        
        # 解析结构化内容
        content_data = json.loads(structured_content)
        elements = content_data.get("elements", [])
        
        if not elements:
            return JSONResponse(
                status_code=400,
                content={"success": False, "error": "内容为空"}
            )
        
        progress_queue.put({
            "type": "step",
            "message": f"[INFO] 解析到 {len(elements)} 个内容块",
            "timestamp": time.time()
        })
        
        # 生成输出文件路径
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"output_structured_{timestamp}.docx"
        output_path = Path("uploads") / output_filename
        
        # 直接生成文档（不通过Agent）
        from docx import Document
        
        progress_queue.put({
            "type": "step",
            "message": f"[INFO] 加载模板样式...",
            "timestamp": time.time()
        })
        
        # 加载模板以获取样式
        doc = Document(template_path)
        
        # 清空模板内容
        for element in doc.element.body[:]:
            if element.tag.endswith('p') or element.tag.endswith('tbl'):
                element.getparent().remove(element)
        
        progress_queue.put({
            "type": "step",
            "message": f"[INFO] 写入内容...",
            "timestamp": time.time()
        })
        
        # 添加内容
        for elem in elements:
            style_name = elem.get("style_name", "Normal")
            text = elem.get("text", "")
            
            if text.strip():
                para = doc.add_paragraph(text)
                try:
                    para.style = style_name
                except:
                    para.style = "Normal"
        
        # 保存文档
        doc.save(str(output_path))
        
        progress_queue.put({
            "type": "complete",
            "message": f"[OK] 文档生成成功！\n\n[FILE] [{output_filename}](/api/download/{output_filename})",
            "timestamp": time.time()
        })
        
        return {
            "success": True,
            "output_path": str(output_path),
            "filename": output_filename
        }
    
    except json.JSONDecodeError as e:
        progress_queue.put({
            "type": "error",
            "message": f"[ERROR] JSON解析错误: {str(e)}",
            "timestamp": time.time()
        })
        return JSONResponse(
            status_code=400,
            content={"success": False, "error": f"JSON解析错误: {str(e)}"}
        )
    except Exception as e:
        progress_queue.put({
            "type": "error",
            "message": f"[ERROR] 处理失败: {str(e)}",
            "timestamp": time.time()
        })
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )



@app.get("/api/status")
async def get_status():
    """检查 Agent 状态"""
    return {
        "initialized": agent is not None,
        "ready": agent is not None
    }

@app.get("/api/download/{filename}")
async def download_file(filename: str):
    """
    下载生成的文件
    """
    try:
        upload_dir = Path("uploads")
        file_path = upload_dir / filename
        
        if not file_path.exists():
            return JSONResponse(
                status_code=404,
                content={"success": False, "error": "文件不存在"}
            )
        
        # 返回文件供下载
        return FileResponse(
            path=str(file_path),
            filename=filename,
            media_type='application/octet-stream'
        )
    
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )

@app.get("/api/preview/{file_type}/{filename}")
async def preview_file(file_type: str, filename: str):
    """
    预览文件
    返回文件内容用于前端预览
    """
    try:
        upload_dir = Path("uploads")
        file_path = upload_dir / filename
        
        if not file_path.exists():
            return JSONResponse(
                status_code=404,
                content={"success": False, "error": "文件不存在"}
            )
        
        # 根据文件类型返回不同内容
        if file_type == "text" or file_type == "markdown":
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return {"success": True, "content": content, "type": file_type}
        
        elif file_type == "docx":
            # 简单提取 Word 文档文本
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return {
                "success": True,
                "content": "\n\n".join(paragraphs),
                "type": "docx"
            }
        
        else:
            return JSONResponse(
                status_code=400,
                content={"success": False, "error": "不支持的预览类型"}
            )
    
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )

if __name__ == "__main__":
    # 在 8765 端口启动（避免与常见端口冲突）
    uvicorn.run(
        app,
        host="127.0.0.1",
        port=8765,
        log_level="info"
    )
